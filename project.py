# # import streamlit as st
# # import requests

# # # Function to handle API requests
# # def call_api(data):
# #     # Replace with your API URL
# #     api_url = "https://llama-1.onrender.com/history"
# #     response = requests.post(api_url, json=data)
# #     if response.status_code == 200:
# #         return response.json()
# #     else:
# #         st.error("API request failed.")
# #         return None

# # # Common input form
# # def input_form():
# #     st.text("Input Form")
    
# #     # Text input
# #     text = st.text_area("Enter your text here:")
    
# #     # PDF Upload
# #     uploaded_file = st.file_uploader("Upload a PDF file", type=["pdf"])
    
# #     # PDF URL input
# #     pdf_url = st.text_input("Enter the URL of a PDF:")
    
# #     # Process inputs
# #     if st.button("Send"):
# #         data = {}
        
# #         if text:
# #             data["input_type"] = "text"
# #             data["content"] = text
# #         elif uploaded_file:
# #             # Convert PDF to binary
# #             files = {"file": uploaded_file.getvalue()}
# #             response = requests.post("https://llama-1.onrender.com/history", files=files)
# #             if response.status_code == 200:
# #                 st.json(response.json())
# #             else:
# #                 st.error("API request failed.")
# #             return
# #         elif pdf_url:
# #             data["input_type"] = "url"
# #             data["content"] = pdf_url
# #         else:
# #             st.error("Please provide text, upload a PDF, or enter a URL.")
# #             return
        
# #         response = call_api(data)
# #         if response:
# #             st.json(response)

# # # Main page
# # def main_page():
# #     st.title("Main Page")
# #     input_form()

# # # Additional pages with similar content
# # def page_two():
# #     st.title("Page Two")
# #     input_form()

# # def page_three():
# #     st.title("Page Three")
# #     input_form()

# # # Sidebar for page navigation
# # st.sidebar.title("Navigation")
# # page = st.sidebar.radio("Go to", ["Main Page", "Page Two", "Page Three"])

# # # Load the selected page
# # if page == "Main Page":
# #     main_page()
# # elif page == "Page Two":
# #     page_two()
# # elif page == "Page Three":
# #     page_three()

# # import streamlit as st
# # import requests

# # # Initialize session state
# # if 'input_data' not in st.session_state:
# #     st.session_state.input_data = {
# #         "text": "",
# #         "pdf_url": ""
# #     }


# # # Function to display API response
# # def display_response(response):
# #     st.subheader("API Response")
    
# #     if isinstance(response, dict):
# #         # Example: If the response contains "answer" and "info" keys
# #         if "answer" in response:
# #             st.write(f"**Answer:** {response['answer']}")
# #         if "info" in response:
# #             st.write(f"**Additional Info:** {response['info']}")
        
# #         # Display other keys
# #         for key, value in response.items():
# #             if key not in ["answer", "info"]:
# #                 st.write(f"**{key.capitalize()}:** {value}")
# #     else:
# #         st.write(response)

# # # Function to handle API requests
# # def call_api(data):
# #     # Replace with your API URL
# #     api_url = "https://llama-1.onrender.com/history"
    
# #     # Constructing the request payload
# #     payload = {
# #         "username": "siddharamsutar23@gmail.com",  # Add your username
# #         "prompt": data["content"]  # Add the prompt
# #     }
    
# #     response = requests.post(api_url, json=payload)
# #     if response.status_code == 200:
# #         return response.json()
# #     else:
# #         st.error("API request failed.")
# #         return None

# # # Common input form
# # def input_form():
# #     st.text("Input Form")
    
# #     # Text input
# #     st.session_state.input_data['text'] = st.text_area("Enter your text here:", value=st.session_state.input_data['text'])
    
# #     # PDF URL input
# #     st.session_state.input_data['pdf_url'] = st.text_input("Enter the URL of a PDF:", value=st.session_state.input_data['pdf_url'])
    
# #     # Process inputs
# #     if st.button("Send"):
# #         data = {}

# #         if st.session_state.input_data['text']:
# #             data["content"] = st.session_state.input_data['text']
# #         elif st.session_state.input_data['pdf_url']:
# #             data["content"] = st.session_state.input_data['pdf_url']
# #         else:
# #             st.error("Please provide text or enter a URL.")
# #             return
        
# #         response = call_api(data)
# #         if response:
# #             display_response(response)



# # # Main page
# # def main_page():
# #     st.title("Main Page")
# #     input_form()

# # # Additional pages with similar content
# # def page_two():
# #     st.title("Page Two")
# #     input_form()

# # def page_three():
# #     st.title("Page Three")
# #     input_form()

# # # Sidebar for page navigation
# # st.sidebar.title("Navigation")
# # page = st.sidebar.radio("Go to", ["Main Page", "Page Two", "Page Three"])

# # # Load the selected page
# # if page == "Main Page":
# #     main_page()
# # elif page == "Page Two":
# #     page_two()
# # elif page == "Page Three":
# #     page_three()
# # import streamlit as st
# # import requests
# # import PyPDF2

# # # Initialize session state
# # if 'input_data' not in st.session_state:
# #     st.session_state.input_data = {
# #         "text": "",
# #         "pdf_url": "",
# #         "pdf_files": None
# #     }

# # # Function to extract text from a list of PDF files
# # def extract_text_from_pdfs(pdf_files):
# #     combined_text = ""
# #     for pdf_file in pdf_files:
# #         try:
# #             reader = PyPDF2.PdfReader(pdf_file)
# #             text = ""
# #             for page_num in range(len(reader.pages)):
# #                 page = reader.pages[page_num]
# #                 text += page.extract_text()
# #             combined_text += text + "\n"  # Separate content from different PDFs
# #         except Exception as e:
# #             st.error(f"Failed to read PDF file: {e}")
# #             return None
# #     return combined_text

# # # Function to handle API requests
# # def call_api(data):
# #     api_url = "https://llama-1.onrender.com/history"

# #     # Construct the JSON payload
# #     payload = {
# #         "username": "siddharamsutar23@gmail.com",
# #         "prompt": data["content"]
# #     }
    
# #     response = requests.post(api_url, json=payload)
    
# #     if response.status_code == 200:
# #         return response.json()
# #     else:
# #         st.error(f"API request failed with status code {response.status_code}.")
# #         return None

# # # Function to display API response
# # def display_response(response):
# #     st.subheader("API Response")

# #     if isinstance(response, dict):
# #         for key, value in response.items():
# #             st.write(f"**{key.capitalize()}:** {value}")
# #     else:
# #         st.write(response)

# # # Common input form
# # def input_form():
# #     st.text("Input Form")

# #     # Text input
# #     st.session_state.input_data['text'] = st.text_area("Enter your text here:", value=st.session_state.input_data['text'])

# #     # PDF URL input
# #     st.session_state.input_data['pdf_url'] = st.text_input("Enter the URL of a PDF:", value=st.session_state.input_data['pdf_url'])

# #     # File upload
# #     st.session_state.input_data['pdf_files'] = st.file_uploader("Upload PDF files", type=["pdf"], accept_multiple_files=True)

# #     # Process inputs
# #     if st.button("Send"):
# #         data = {}

# #         if st.session_state.input_data['text']:
# #             data["content"] = st.session_state.input_data['text']
# #         elif st.session_state.input_data['pdf_url']:
# #             data["content"] = st.session_state.input_data['pdf_url']
# #         elif st.session_state.input_data['pdf_files']:
# #             try:
# #                 pdf_text = extract_text_from_pdfs(st.session_state.input_data['pdf_files'])
# #                 if pdf_text:
# #                     data["content"] = pdf_text
# #                 else:
# #                     return
# #             except Exception as e:
# #                 st.error(f"Failed to extract text from PDFs: {e}")
# #                 return
# #         else:
# #             st.error("Please provide text, upload a PDF, or enter a URL.")
# #             return

# #         response = call_api(data)
# #         if response:
# #             display_response(response)

# # # Main page
# # def main_page():
# #     st.title("Main Page")
# #     input_form()

# # # Additional pages with similar content
# # def page_two():
# #     st.title("Page Two")
# #     input_form()

# # def page_three():
# #     st.title("Page Three")
# #     input_form()

# # # Sidebar for page navigation
# # st.sidebar.title("Navigation")
# # page = st.sidebar.radio("Go to", ["Main Page", "Page Two", "Page Three"])

# # # Load the selected page
# # if page == "Main Page":
# #     main_page()
# # elif page == "Page Two":
# #     page_two()
# # elif page == "Page Three":
# #     page_three()

import streamlit as st
import requests
import PyPDF2
from PIL import Image
import docx

# Initialize session state
if 'input_data' not in st.session_state:
    st.session_state.input_data = {
        "text": "",
        "pdf_url": "",
        "files": None
    }

# Function to extract text from PDFs
def extract_text_from_pdfs(pdf_files):
    combined_text = ""
    for pdf_file in pdf_files:
        try:
            reader = PyPDF2.PdfReader(pdf_file)
            text = ""
            for page_num in range(len(reader.pages)):
                page = reader.pages[page_num]
                text += page.extract_text()
            combined_text += text + "\n"
        except Exception as e:
            st.error(f"Failed to read PDF file: {e}")
            return None
    return combined_text

# Function to extract text from DOCX files
def extract_text_from_docx(docx_files):
    combined_text = ""
    for docx_file in docx_files:
        try:
            doc = docx.Document(docx_file)
            text = "\n".join([para.text for para in doc.paragraphs])
            combined_text += text + "\n"
        except Exception as e:
            st.error(f"Failed to read DOCX file: {e}")
            return None
    return combined_text

# Function to process image files (e.g., PNG, JPG)
def process_images(image_files):
    image_texts = []
    for image_file in image_files:
        try:
            img = Image.open(image_file)
            st.image(img, caption=image_file.name)
            image_texts.append(f"Image file: {image_file.name}")
        except Exception as e:
            st.error(f"Failed to process image: {e}")
    return "\n".join(image_texts)

# Function to handle API requests
def call_api(data):
    api_url = "https://llama-1.onrender.com/history"

    # Construct the JSON payload
    payload = {
        "username": "siddharamsutar23@gmail.com",
        "prompt": data["content"]
    }
    
    response = requests.post(api_url, json=payload)
    
    if response.status_code == 200:
        return response.json()
    else:
        st.error(f"API request failed with status code {response.status_code}.")
        return None
    
def call_api1(data):
    api_url = "https://mistralai-axq4.onrender.com/history"

    # Construct the JSON payload
    payload = {
        "username": "siddharamsutar23@gmail.com",
        "prompt": data["content"]
    }
    
    response = requests.post(api_url, json=payload)
    
    if response.status_code == 200:
        return response.json()
    else:
        st.error(f"API request failed with status code {response.status_code}.")
        return None
    
    
def call_api2(data):
    api_url = "https://jsgemiintegration.onrender.com/history"

    # Construct the JSON payload
    payload = {
        "username": "siddharamsutar23@gmail.com",
        "prompt": data["content"]
    }
    
    response = requests.post(api_url, json=payload)
    
    if response.status_code == 200:
        return response.json()
    else:
        st.error(f"API request failed with status code {response.status_code}.")
        return None


# Function to display API response
def display_response(response):
    st.subheader("API Response")
    if isinstance(response, dict):
        for key, value in response.items():
            st.write(f"**{key.capitalize()}:** {value}")
    else:
        st.write(response)

# Common input form
def input_form():
    st.text("Input Form")

    # Text input
    st.session_state.input_data['text'] = st.text_area("Enter your text here:", value=st.session_state.input_data['text'])

    # PDF URL input
    st.session_state.input_data['pdf_url'] = st.text_input("Enter the URL of a PDF:", value=st.session_state.input_data['pdf_url'])

    # File upload
    st.session_state.input_data['files'] = st.file_uploader(
        "Upload files (PDF, DOCX, PNG, JPG)",
        type=["pdf", "docx", "png", "jpg", "jpeg"],
        accept_multiple_files=True
    )

    # Process inputs
    if st.button("Send"):
        data = {}

        if st.session_state.input_data['text']:
            data["content"] = st.session_state.input_data['text']
        elif st.session_state.input_data['pdf_url']:
            data["content"] = st.session_state.input_data['pdf_url']
        elif st.session_state.input_data['files']:
            files = st.session_state.input_data['files']
            combined_text = ""

            pdf_files = [f for f in files if f.name.endswith('.pdf')]
            if pdf_files:
                pdf_text = extract_text_from_pdfs(pdf_files)
                if pdf_text:
                    combined_text += pdf_text

            docx_files = [f for f in files if f.name.endswith('.docx')]
            if docx_files:
                docx_text = extract_text_from_docx(docx_files)
                if docx_text:
                    combined_text += docx_text

            image_files = [f for f in files if f.name.lower().endswith(('.png', '.jpg', '.jpeg'))]
            if image_files:
                image_text = process_images(image_files)
                combined_text += image_text

            if combined_text:
                data["content"] = combined_text
            else:
                st.error("No valid content to process.")
                return
        else:
            st.error("Please provide text, upload files, or enter a URL.")
            return

        response = call_api(data)
        if response:
            display_response(response)


def input_form1():
    st.text("Input Form")

    # Text input
    st.session_state.input_data['text'] = st.text_area("Enter your text here:", value=st.session_state.input_data['text'])

    # PDF URL input
    st.session_state.input_data['pdf_url'] = st.text_input("Enter the URL of a PDF:", value=st.session_state.input_data['pdf_url'])

    # File upload
    st.session_state.input_data['files'] = st.file_uploader(
        "Upload files (PDF, DOCX, PNG, JPG)",
        type=["pdf", "docx", "png", "jpg", "jpeg"],
        accept_multiple_files=True
    )

    # Process inputs
    if st.button("Send"):
        data = {}

        if st.session_state.input_data['text']:
            data["content"] = st.session_state.input_data['text']
        elif st.session_state.input_data['pdf_url']:
            data["content"] = st.session_state.input_data['pdf_url']
        elif st.session_state.input_data['files']:
            files = st.session_state.input_data['files']
            combined_text = ""

            pdf_files = [f for f in files if f.name.endswith('.pdf')]
            if pdf_files:
                pdf_text = extract_text_from_pdfs(pdf_files)
                if pdf_text:
                    combined_text += pdf_text

            docx_files = [f for f in files if f.name.endswith('.docx')]
            if docx_files:
                docx_text = extract_text_from_docx(docx_files)
                if docx_text:
                    combined_text += docx_text

            image_files = [f for f in files if f.name.lower().endswith(('.png', '.jpg', '.jpeg'))]
            if image_files:
                image_text = process_images(image_files)
                combined_text += image_text

            if combined_text:
                data["content"] = combined_text
            else:
                st.error("No valid content to process.")
                return
        else:
            st.error("Please provide text, upload files, or enter a URL.")
            return

        response = call_api1(data)
        if response:
            display_response(response)



def input_form2():
    st.text("Input Form")

    # Text input
    st.session_state.input_data['text'] = st.text_area("Enter your text here:", value=st.session_state.input_data['text'])

    # PDF URL input
    st.session_state.input_data['pdf_url'] = st.text_input("Enter the URL of a PDF:", value=st.session_state.input_data['pdf_url'])

    # File upload
    st.session_state.input_data['files'] = st.file_uploader(
        "Upload files (PDF, DOCX, PNG, JPG)",
        type=["pdf", "docx", "png", "jpg", "jpeg"],
        accept_multiple_files=True
    )

    # Process inputs
    if st.button("Send"):
        data = {}

        if st.session_state.input_data['text']:
            data["content"] = st.session_state.input_data['text']
        elif st.session_state.input_data['pdf_url']:
            data["content"] = st.session_state.input_data['pdf_url']
        elif st.session_state.input_data['files']:
            files = st.session_state.input_data['files']
            combined_text = ""

            pdf_files = [f for f in files if f.name.endswith('.pdf')]
            if pdf_files:
                pdf_text = extract_text_from_pdfs(pdf_files)
                if pdf_text:
                    combined_text += pdf_text

            docx_files = [f for f in files if f.name.endswith('.docx')]
            if docx_files:
                docx_text = extract_text_from_docx(docx_files)
                if docx_text:
                    combined_text += docx_text

            image_files = [f for f in files if f.name.lower().endswith(('.png', '.jpg', '.jpeg'))]
            if image_files:
                image_text = process_images(image_files)
                combined_text += image_text

            if combined_text:
                data["content"] = combined_text
            else:
                st.error("No valid content to process.")
                return
        else:
            st.error("Please provide text, upload files, or enter a URL.")
            return

        response = call_api2(data)
        if response:
            display_response(response)

# Main page
def main_page():
    st.title("Main Page")
     # Apply custom CSS styles (Optional, if you want the styles on all pages)
    st.markdown(
        """
        <style>
        /* Sidebar background */
        [data-testid="stSidebar"] {
            background-color: black;
        }
        /* Sidebar text color */
        [data-testid="stSidebar"] * {
            color: white;
        }
        /* Upload button container background */
        .stFileUpload div[role="button"] {
            background-color: #FFF8E1; /* Cream yellow */
            color: black;
            border: 1px solid #DDD;
            border-radius: 4px;
        }
        /* Drag and drop area background */
        .stFileUpload .stDropzone {
            background-color: #FFF8E1;
            border: 1px dashed #DDD;
            border-radius: 4px;
        }
        /* Submit & Process button background */
        .stButton > button {
            background-color: #FFF8E1;
            color: black;
            border: 1px solid #DDD;
            border-radius: 4px;
        }
        </style>
        """,
        unsafe_allow_html=True
    )
    input_form()

# Additional pages with similar content
def page_two():
    st.title("Page Two")
     # Apply custom CSS styles (Optional, if you want the styles on all pages)
    st.markdown(
        """
        <style>
        /* Sidebar background */
        [data-testid="stSidebar"] {
            background-color: black;
        }
        /* Sidebar text color */
        [data-testid="stSidebar"] * {
            color: white;
        }
        /* Upload button container background */
        .stFileUpload div[role="button"] {
            background-color: #FFF8E1; /* Cream yellow */
            color: black;
            border: 1px solid #DDD;
            border-radius: 4px;
        }
        /* Drag and drop area background */
        .stFileUpload .stDropzone {
            background-color: #FFF8E1;
            border: 1px dashed #DDD;
            border-radius: 4px;
        }
        /* Submit & Process button background */
        .stButton > button {
            background-color: #FFF8E1;
            color: black;
            border: 1px solid #DDD;
            border-radius: 4px;
        }
        </style>
        """,
        unsafe_allow_html=True
    )
    input_form1()

def page_three():
    st.title("Page Three")
     # Apply custom CSS styles (Optional, if you want the styles on all pages)
    st.markdown(
        """
        <style>
        /* Sidebar background */
        [data-testid="stSidebar"] {
            background-color: black;
        }
        /* Sidebar text color */
        [data-testid="stSidebar"] * {
            color: white;
        }
        /* Upload button container background */
        .stFileUpload div[role="button"] {
            background-color: #FFF8E1; /* Cream yellow */
            color: black;
            border: 1px solid #DDD;
            border-radius: 4px;
        }
        /* Drag and drop area background */
        .stFileUpload .stDropzone {
            background-color: #FFF8E1;
            border: 1px dashed #DDD;
            border-radius: 4px;
        }
        /* Submit & Process button background */
        .stButton > button {
            background-color: #FFF8E1;
            color: black;
            border: 1px solid #DDD;
            border-radius: 4px;
        }
        </style>
        """,
        unsafe_allow_html=True
    )
    input_form2()

# Sidebar for page navigation
st.sidebar.title("Navigation")
page = st.sidebar.radio("Go to", ["Main Page", "Page Two", "Page Three"])

# Load the selected page
if page == "Main Page":
    main_page()
elif page == "Page Two":
    page_two()
elif page == "Page Three":
    page_three()
#########################################
# import streamlit as st
# import requests
# import PyPDF2
# from PIL import Image
# # import docx
# import streamlit as st
# import requests
# import PyPDF2
# from PIL import Image
# import docx
# import io

# # Initialize session state
# if 'input_data' not in st.session_state:
#     st.session_state.input_data = {
#         "text": "",
#         "pdf_url": "",
#         "files": None
#     }

# Function to simulate file upload by reading files from the disk
def programmatically_upload_files(file_paths):
    uploaded_files = []
    for file_path in file_paths:
        with open(file_path, "rb") as f:
            pdf_bytes = f.read()
            file_obj = io.BytesIO(pdf_bytes)
            file_obj.name = file_path  # Set the name attribute manually
            uploaded_files.append(file_obj)
    st.session_state.input_data['files'] = uploaded_files

# Function to extract text from PDFs
def extract_text_from_pdfs(pdf_files):
    combined_text = ""
    for pdf_file in pdf_files:
        try:
            reader = PyPDF2.PdfReader(pdf_file)
            text = ""
            for page_num in range(len(reader.pages)):
                page = reader.pages[page_num]
                text += page.extract_text()
            combined_text += text + "\n"
        except Exception as e:
            st.error(f"Failed to read PDF file: {e}")
            return None
    return combined_text

# Function to extract text from DOCX files
def extract_text_from_docx(docx_files):
    combined_text = ""
    for docx_file in docx_files:
        try:
            doc = docx.Document(docx_file)
            text = "\n".join([para.text for para in doc.paragraphs])
            combined_text += text + "\n"
        except Exception as e:
            st.error(f"Failed to read DOCX file: {e}")
            return None
    return combined_text

# Function to process image files (e.g., PNG, JPG)
def process_images(image_files):
    image_texts = []
    for image_file in image_files:
        try:
            img = Image.open(image_file)
            st.image(img, caption=image_file.name)
            image_texts.append(f"Image file: {image_file.name}")
        except Exception as e:
            st.error(f"Failed to process image: {e}")
    return "\n".join(image_texts)

# Function to handle API requests
def call_api(data):
    api_url = "https://llama-1.onrender.com/history"

    # Construct the JSON payload
    payload = {
        "username": "siddharamsutar23@gmail.com",
        "prompt": data["content"]
    }
    
    response = requests.post(api_url, json=payload)
    
    if response.status_code == 200:
        return response.json()
    else:
        st.error(f"API request failed with status code {response.status_code}.")
        return None

# Function to display API response
def display_response(response):
    st.subheader("API Response")
    if isinstance(response, dict):
        for key, value in response.items():
            st.write(f"**{key.capitalize()}:** {value}")
    else:
        st.write(response)

# Main page
def main_page():
    st.title("Main Page")

    # Programmatically upload files
    programmatically_upload_files([
        r"C:\Users\HP\Downloads\Doj1_merged.pdf"

    ])

    # Process the uploaded files
    if st.session_state.input_data['files']:
        files = st.session_state.input_data['files']
        combined_text = ""

        # Process PDF files
        pdf_files = [f for f in files if f.name.endswith('.pdf')]
        if pdf_files:
            pdf_text = extract_text_from_pdfs(pdf_files)
            if pdf_text:
                combined_text += pdf_text

        # Process DOCX files
        docx_files = [f for f in files if f.name.endswith('.docx')]
        if docx_files:
            docx_text = extract_text_from_docx(docx_files)
            if docx_text:
                combined_text += docx_text

        # Process Image files
        image_files = [f for f in files if f.name.lower().endswith(('.png', '.jpg', '.jpeg'))]
        if image_files:
            image_text = process_images(image_files)
            combined_text += image_text

        if combined_text:
            prompt_text = st.text_input("Enter your question or prompt:")

            if st.button("Generate Response"):
                # Combine prompt with extracted text
                full_prompt = f"{prompt_text}\n\n{combined_text}"
                
                data = {"content": full_prompt}
                
                # Call the API with the combined content
                api_response = call_api(data)
                
                if api_response:
                    st.subheader("AI Response")
                    st.write(api_response)


        else:
            st.error("No valid content to process.")
    else:
        st.error("No files found to process.")

# Run the app
if __name__ == "__main__":
    main_page()
import streamlit as st
import requests
import PyPDF2
from PIL import Image
import docx
import io
import json
import google.generativeai as genai
import os

# Initialize session state
if 'input_data' not in st.session_state:
    st.session_state.input_data = {
        "text": "",
        "pdf_url": "",
        "files": None,
        "json_prompt": ""  # Field for JSON input prompt
    }

# # Function to simulate file upload by reading files from the disk
# def programmatically_upload_files(file_paths):
#     uploaded_files = []
#     for file_path in file_paths:
#         with open(file_path, "rb") as f:
#             file_bytes = f.read()
#             file_obj = io.BytesIO(file_bytes)
#             file_obj.name = file_path  # Set the name attribute manually
#             uploaded_files.append(file_obj)
#     st.session_state.input_data['files'] = uploaded_files

# # Function to extract text from PDFs
# def extract_text_from_pdfs(pdf_files):
#     combined_text = ""
#     for pdf_file in pdf_files:
#         try:
#             reader = PyPDF2.PdfReader(pdf_file)
#             text = ""
#             for page_num in range(len(reader.pages)):
#                 page = reader.pages[page_num]
#                 text += page.extract_text()
#             combined_text += text + "\n"
#         except Exception as e:
#             st.error(f"Failed to read PDF file: {e}")
#             return None
#     return combined_text

# # Function to extract text from DOCX files
# def extract_text_from_docx(docx_files):
#     combined_text = ""
#     for docx_file in docx_files:
#         try:
#             doc = docx.Document(docx_file)
#             text = "\n".join([para.text for para in doc.paragraphs])
#             combined_text += text + "\n"
#         except Exception as e:
#             st.error(f"Failed to read DOCX file: {e}")
#             return None
#     return combined_text

# # Function to process image files (e.g., PNG, JPG)
# def process_images(image_files):
#     image_texts = []
#     for image_file in image_files:
#         try:
#             img = Image.open(image_file)
#             image_texts.append(f"Image file: {image_file.name}")
#         except Exception as e:
#             st.error(f"Failed to process image: {e}")
#     return "\n".join(image_texts)
# genai.configure(api_key=os.getenv("GOOGLE_API_KEY"))
# model = genai.GenerativeModel("gemini-1.5-flash")
# # # Function to handle API requests
# def generate_response(data):
#     prompt=data['content']

#     response = model.generate_content(prompt)
#     if response:
#         return response.json()
#     else:
#         st.error("errror")



# # Function to display API response
# def display_response(response):
#     st.subheader("AI Response")
#     if isinstance(response, dict):
#         for key, value in response.items():
#             st.write(f"**{key.capitalize()}:** {value}")
#     else:
#         st.write(response)

# # Main page
# def main_page():
#     st.title("Main Page")

#     # Programmatically upload files
#     programmatically_upload_files([
#         "C:/Users/HP/Downloads/Tele Law.pdf",
#     ])

#     # JSON input field
#     json_input = st.text_area("Enter JSON input:")
#     if json_input:
#         try:
#             json_data = json.loads(json_input)
#             if 'prompt' in json_data:
#                 st.session_state.input_data['json_prompt'] = json_data['prompt']
#             else:
#                 st.error("JSON input must contain a 'prompt' field.")
#                 return
#         except json.JSONDecodeError:
#             st.error("Invalid JSON format.")
#             return

#     # Process the uploaded files
#     if st.session_state.input_data['files']:
#         files = st.session_state.input_data['files']
#         combined_text = ""

#         # Process PDF files
#         pdf_files = [f for f in files if f.name.endswith('.pdf')]
#         if pdf_files:
#             pdf_text = extract_text_from_pdfs(pdf_files)
#             if pdf_text:
#                 combined_text += pdf_text

#         # Process DOCX files
#         docx_files = [f for f in files if f.name.endswith('.docx')]
#         if docx_files:
#             docx_text = extract_text_from_docx(docx_files)
#             if docx_text:
#                 combined_text += docx_text

#         # Process Image files
#         image_files = [f for f in files if f.name.lower().endswith(('.png', '.jpg', '.jpeg'))]
#         if image_files:
#             image_text = process_images(image_files)
#             combined_text += image_text

#         if combined_text or st.session_state.input_data['json_prompt']:
#             prompt_text = st.text_input("Enter your question or prompt:")

#             if st.button("Generate Response"):
#                 # Combine prompt with extracted text and JSON input prompt
#                 full_prompt = f"{prompt_text}\n\n{combined_text}"
                
#                 if st.session_state.input_data['json_prompt']:
#                     full_prompt += f"\n\n{st.session_state.input_data['json_prompt']}"
                
#                 data = {"content": full_prompt}
                
#                 # Call the API with the combined content
#                 api_response = generate_response(data)
                
#                 if api_response:
#                     st.subheader("AI Response")
#                     st.write(api_response)

#         else:
#             st.error("No valid content to process.")
#     else:
#         st.error("No files found to process.")

# # Run the app
# if __name__ == "__main__":
#     main_page()
